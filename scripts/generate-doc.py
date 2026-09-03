import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_user_guide():
    doc = Document()

    # Configure 1-inch margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Color Palette: Deep Blue / Teal / Charcoal
    PRIMARY_COLOR = RGBColor(27, 85, 226)    # #1B55E2
    SECONDARY_COLOR = RGBColor(86, 212, 221) # #56D4DD
    DARK_TEXT = RGBColor(30, 41, 59)        # #1E293B
    MUTED_TEXT = RGBColor(100, 116, 139)    # #64748B

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("Proof of Demand (PoD) Engine")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(28)
    r_title.font.bold = True
    r_title.font.color.rgb = PRIMARY_COLOR

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("The Complete Founder's Guide: Validating Startup Demand Before Building")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = MUTED_TEXT

    # Metadata
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Version 1.0  |  Platform: https://pod-blue-nine.vercel.app  |  September 2026")
    r_meta.font.name = "Arial"
    r_meta.font.size = Pt(10)
    r_meta.font.italic = True
    r_meta.font.color.rgb = MUTED_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for Section Headings
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_COLOR
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = DARK_TEXT
        return h

    def add_body(text, bold_prefix=None, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Arial"
            rb.font.size = Pt(10.5)
            rb.font.bold = True
            rb.font.color.rgb = DARK_TEXT
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK_TEXT
        return p

    def add_callout(text, title="KEY TAKEAWAY"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "F0F7FF") # light ice blue
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_head = p.add_run(f"💡 {title}: ")
        r_head.font.name = "Arial"
        r_head.font.size = Pt(10)
        r_head.font.bold = True
        r_head.font.color.rgb = PRIMARY_COLOR
        
        r_text = p.add_run(text)
        r_text.font.name = "Arial"
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = DARK_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. Executive Summary
    add_heading_1("1. What is PoD Engine?")
    add_body(
        "Proof of Demand (PoD) Engine is an evidence-first validation platform for founders, product builders, "
        "and startup studios. More than 90% of new startups fail because they spend months writing code for products "
        "that customers ultimately do not want or will not pay for. PoD Engine flips this traditional approach on its head: "
        "it validates market demand using real behavioral evidence, live micro-landing pages, and AI-powered intent scoring "
        "BEFORE you build."
    )
    add_callout(
        "Never spend 6 months building an MVP on a guess. Use PoD Engine to test positioning, audience resonance, "
        "and willingness to pay in 3 to 7 days for under $100.",
        "THE FOUNDER'S GOLDEN RULE"
    )

    # 2. Who Uses This App?
    add_heading_1("2. Who Uses This Platform?")
    add_body(" Solo founders testing a new software idea before quitting their job or hiring developers.", "• Solo Founders & Indie Hackers: ")
    add_body(" Incubators, accelerators, and venture builders managing a portfolio of 10+ concurrent ideas to determine which concepts deserve funding.", "• Startup Studios & Venture Builders: ")
    add_body(" Evaluating feature expansions, pricing tier adjustments, or entering adjacent market segments.", "• Growth & Product Teams: ")

    # 3. How Someone Uses the App (End-to-End Workflow)
    add_heading_1("3. End-to-End User Journey (How to Use the App)")

    add_heading_2("Step 1: Sign In & Workspace Setup")
    add_body(
        "Users sign in securely via Clerk using their Google account or email. Upon authentication, each user receives a "
        "dedicated workspace. If you work with co-founders or team members, you can invite them by email under Settings > Team, "
        "assigning roles such as Owner, Admin, or Member."
    )

    add_heading_2("Step 2: The Idea & Onboarding Wizard (/onboarding)")
    add_body(
        "When starting a new venture, the user enters the 6-step Onboarding Wizard to define the core assumptions:"
    )
    add_body(" The name and a crisp one-sentence description of what your solution does.", "1. Product Concept: ")
    add_body(" What manual process or frustration currently costs your target customer hours or dollars?", "2. Core Problem: ")
    add_body(" How do customers solve this today? (e.g. spreadsheets, competitor software, manual labor)", "3. Alternatives: ")
    add_body(" Subscription, usage-based, or one-time license, including expected price points ($49/mo, $99/mo, etc.).", "4. Business Model & Pricing: ")
    add_body(" Who makes the buying decision? (Job titles, industry, geography, company headcount).", "5. Target Audience: ")
    add_body(" Expected test ad spend ($50 - $200) across Meta, Google, or LinkedIn.", "6. Validation Budget: ")

    add_heading_2("Step 3: Audience Personas & Segments (/dashboard/audiences)")
    add_body(
        "The Audiences dashboard tracks your primary customer profile and breaks them down into segmented cohorts. "
        "For example, a B2B SaaS tool might segment between 'Early Adopter Founders' (high agility, rapid conversion) "
        "and 'Growth Stage Operations Leads' (higher budget, longer buying cycle). The app calculates intent scores for each segment."
    )

    add_heading_2("Step 4: Launching Landing Page Variants (/dashboard/landing-pages)")
    add_body(
        "Testing a single message is rarely enough. The app allows you to launch multiple live landing page variants "
        "hosted instantly at https://pod-blue-nine.vercel.app/p/[slug]. You can test contrasting angles:"
    )
    add_body(" Focuses on the acute pain point ('Stop wasting 6 hours every Friday').", "• Variant A (Pain-Centric): ")
    add_body(" Focuses on the technological capability ('AI that automates report generation in 1 click').", "• Variant B (Feature-Centric): ")
    add_body(" Tests transparent pricing tiers to gauge purchase readiness before building.", "• Variant C (Pricing-Centric): ")

    add_heading_2("Step 5: Tracking Real Behavioral Demand Signals (/dashboard/signals)")
    add_body(
        "Traditional surveys lie; user behavior does not. The PoD Engine tracking pixel automatically monitors high-intent actions "
        "on your landing pages, categorizing them into a structured conversion funnel:"
    )
    add_body(" Total unique visitors landed on the page.", "1. Page View (Awareness): ")
    add_body(" Visitor read past the fold (verified engagement).", "2. Scroll Depth (Interest): ")
    add_body(" Visitor clicked primary value proposition buttons.", "3. CTA Click (Intent): ")
    add_body(" Visitor inspected pricing options and packaging.", "4. Pricing View (Consideration): ")
    add_body(" Visitor submitted their email or clicked 'Order / Pre-order' (Commitment).", "5. Checkout / Form Submit: ")

    add_heading_2("Step 6: Understanding the PoD Demand Score (/dashboard)")
    add_body(
        "The Proof of Demand Score is an algorithmic rating from 0 to 100 derived from visitor volume, high-intent actions, "
        "and willingness-to-pay signals. Based on your score, the system delivers an actionable verdict:"
    )

    # Verdicts Table
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["PoD Score", "Verdict", "Recommendation", "Next Action"]
    col_widths = [Inches(1.1), Inches(1.4), Inches(1.5), Inches(2.5)]

    # Header formatting
    for col_idx, text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.width = col_widths[col_idx]
        set_cell_background(cell, "1B55E2")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    verdict_data = [
        ("75 – 100", "Strong Demand", "BUILD (Green)", "Proceed to development. High conviction of customer willingness to pay."),
        ("50 – 74", "Promising", "ITERATE (Blue)", "Signals are positive. Refine positioning or test an adjusted price point."),
        ("30 – 49", "Needs Iteration", "PAUSE (Amber)", "Weak resonance. Customer problem is not acute enough. Consider pivot."),
        ("< 30", "Weak Demand", "KILL (Red)", "Save your capital. Evidence indicates minimal commercial interest."),
    ]

    for row_idx, row_vals in enumerate(verdict_data, start=1):
        for col_idx, val in enumerate(row_vals):
            cell = table.cell(row_idx, col_idx)
            cell.width = col_widths[col_idx]
            bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Arial"
            r.font.size = Pt(9)
            if col_idx == 1:
                r.font.bold = True
            r.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_heading_2("Step 7: Conversing with the AI Analyst (/dashboard/ai-analyst)")
    add_body(
        "Instead of parsing raw spreadsheets, users can chat directly with the AI Analyst. The AI Analyst inspects "
        "real visitor metrics in your database to answer critical strategic questions:"
    )
    add_body(" 'Which landing page headline generated the lowest bounce rate and highest pre-order intent?'", "• Example Prompt: ")
    add_body(" 'What is the drop-off rate between pricing page views and email submissions?'", "• Example Prompt: ")
    add_body(" 'Given our current 8.4% conversion rate, what CAC can we sustain at a $99/mo price point?'", "• Example Prompt: ")

    add_heading_2("Step 8: Prototyping & Studio Hand-Off (/dashboard/firstmile)")
    add_body(
        "Once an experiment achieves a 'Build' recommendation (PoD Score ≥ 75), the founder can connect their validated "
        "portfolio directly to FirstMileDevs Studio for rapid prototyping, transitioning from demand proof to a functional product in weeks."
    )

    # 4. Summary Checklist
    add_heading_1("4. Founder's Quick-Start Checklist")
    add_body(" Sign up at https://pod-blue-nine.vercel.app/sign-up and configure your workspace.", "[ ] Step 1: ")
    add_body(" Run the Onboarding Wizard to define your product, problem statement, and price.", "[ ] Step 2: ")
    add_body(" Generate at least 2 distinct landing page variants with contrasting value propositions.", "[ ] Step 3: ")
    add_body(" Drive 100 to 500 targeted visitors using Google or Meta ad credits.", "[ ] Step 4: ")
    add_body(" Check your PoD Score: If ≥ 75, begin building; if < 50, iterate or pivot.", "[ ] Step 5: ")
    add_body(" Export executive validation reports for angel investors, co-founders, or grant applications.", "[ ] Step 6: ")

    # Save to disk
    output_path = os.path.abspath("Proof_of_Demand_Engine_User_Guide.docx")
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_user_guide()
